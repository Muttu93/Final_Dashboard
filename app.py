from flask import Flask, render_template, request, redirect, url_for, send_file, flash
import os
from werkzeug.utils import secure_filename
from datetime import datetime
from docx import Document

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # change this if needed

app.config['DATA_FOLDER'] = 'data'
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'gif', 'txt', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    sections = os.listdir(app.config['DATA_FOLDER'])
    return render_template('index.html', sections=sections)

@app.route('/dashboard/<section>')
def dashboard(section):
    section_path = os.path.join(app.config['DATA_FOLDER'], section)
    files = []
    if os.path.exists(section_path):
        for filename in os.listdir(section_path):
            path = os.path.join(section_path, filename)
            if os.path.isfile(path):
                mod_time = datetime.fromtimestamp(os.path.getmtime(path)).strftime('%Y-%m-%d %H:%M:%S')
                files.append({'name': filename, 'modified': mod_time})
    return render_template('index.html', sections=os.listdir(app.config['DATA_FOLDER']),
                           current_section=section, files=files)

@app.route('/upload/<section>', methods=['POST'])
def upload_file(section):
    if 'file' not in request.files:
        flash('No file part!', 'danger')
        return redirect(url_for('dashboard', section=section))
    file = request.files['file']
    if file.filename == '':
        flash('No selected file!', 'warning')
        return redirect(url_for('dashboard', section=section))
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        save_path = os.path.join(app.config['DATA_FOLDER'], section, filename)
        file.save(save_path)
        flash('File uploaded successfully!', 'success')
    return redirect(url_for('dashboard', section=section))

@app.route('/delete/<section>/<filename>')
def delete_file(section, filename):
    file_path = os.path.join(app.config['DATA_FOLDER'], section, filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        flash('File deleted successfully!', 'success')
    else:
        flash('File not found!', 'danger')
    return redirect(url_for('dashboard', section=section))

@app.route('/view/<section>/<filename>')
def view_file(section, filename):
    file_path = os.path.join(app.config['DATA_FOLDER'], section, filename)
    file_ext = os.path.splitext(filename)[1].lower()

    if not os.path.exists(file_path):
        return "File not found", 404

    if file_ext in ['.pdf', '.png', '.jpg', '.jpeg', '.gif']:
        return send_file(file_path)
    elif file_ext in ['.txt', '.docx']:
        return redirect(url_for('edit_file', section=section, filename=filename))
    else:
        flash('Preview not supported for this file type.', 'warning')
        return redirect(url_for('dashboard', section=section))

@app.route('/edit/<section>/<filename>', methods=['GET', 'POST'])
def edit_file(section, filename):
    file_path = os.path.join(app.config['DATA_FOLDER'], section, filename)
    file_ext = os.path.splitext(filename)[1].lower()

    if not os.path.exists(file_path):
        return "File not found", 404

    if request.method == 'POST':
        content = request.form.get('content')
        if file_ext == '.txt':
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        elif file_ext == '.docx':
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)
        flash('Changes saved successfully!', 'success')
        return redirect(url_for('dashboard', section=section))

    content = ''
    if file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    elif file_ext == '.docx':
        doc = Document(file_path)
        content = '\n'.join([para.text for para in doc.paragraphs])

    return render_template('edit.html', filename=filename, content=content, section=section)

@app.route('/search/<section>', methods=['GET'])
def search_file(section):
    query = request.args.get('query', '').lower()
    section_path = os.path.join(app.config['DATA_FOLDER'], section)
    results = []

    if os.path.exists(section_path):
        for filename in os.listdir(section_path):
            if query in filename.lower():
                path = os.path.join(section_path, filename)
                mod_time = datetime.fromtimestamp(os.path.getmtime(path)).strftime('%Y-%m-%d %H:%M:%S')
                results.append({'name': filename, 'modified': mod_time})

    return render_template('index.html', sections=os.listdir(app.config['DATA_FOLDER']),
                           current_section=section, files=results)

if __name__ == '__main__':
    import webbrowser
    webbrowser.open('http://127.0.0.1:5000/')
    app.run(debug=True)
